"""Generate the RecoMart DOCX report from config/report_final.yaml."""
from __future__ import annotations
import argparse, fnmatch, json, re, shutil, sqlite3, subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence
import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from docx.image.image import Image as DocxImage

ROOT=Path(__file__).resolve().parents[2]
DEFAULT_CONFIG=ROOT/"config/report_final.yaml"
BLUE,DARK,MUTED=RGBColor(46,116,181),RGBColor(31,77,120),RGBColor(89,89,89)
WIDTH,INDENT=9360,120
_IMAGE_HASHES=set()
class ReportError(RuntimeError): pass

def load_config(path):
    cfg=yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    for key in ("report","format","team","sections","artifacts","validation"):
        if key not in cfg: raise ReportError(f"Missing configuration section: {key}")
    enabled=[s for s in cfg["sections"] if s.get("enabled",True)]
    ids=[s["id"] for s in enabled]
    if ids!=cfg["validation"].get("required_section_ids",[]): raise ReportError("Enabled section order differs from validation.required_section_ids")
    ownership={"objectives":"assignment","implementation_details":"technical_docs","results":"repository_evidence"}
    by_id={section["id"]:section for section in enabled}
    for section_id,source in ownership.items():
        if by_id.get(section_id,{}).get("source")!=source: raise ReportError(f"{section_id} must use source {source}")
    problem=by_id.get("problem_statement",{})
    forbidden=("sources","source","include_topics","render_mode","prompt")
    if problem.get("type")!="markdown_file" or problem.get("source_file")!="docs/assignment/problem_statement.md":
        raise ReportError("Problem Statement must use only docs/assignment/problem_statement.md")
    if any(key in problem for key in forbidden):
        raise ReportError("Problem Statement contains an additional or prohibited source property")
    assignment=cfg.get("content_sources",{}).get("assignment",[])
    if assignment!=["docs/assignment/ASSIGNMENT_CONTEXT.md"]: raise ReportError("Assignment source must be ASSIGNMENT_CONTEXT.md only")
    return cfg

def field(p,instruction,display=""):
    r=p.add_run(); a=OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"),"begin")
    b=OxmlElement("w:instrText"); b.set(qn("xml:space"),"preserve"); b.text=instruction
    c=OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"),"separate"); d=OxmlElement("w:t"); d.text=display
    e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end"); r._r.extend((a,b,c,d,e))

def setup(doc,cfg,date):
    f,r=cfg["format"],cfg["report"]; sec=doc.sections[0]
    if str(f["page_size"]).upper()!="A4": raise ReportError("Only A4 is supported")
    portrait=str(f["orientation"]).lower()=="portrait"; sec.orientation=WD_ORIENT.PORTRAIT if portrait else WD_ORIENT.LANDSCAPE
    sec.page_width,sec.page_height=((Cm(21),Cm(29.7)) if portrait else (Cm(29.7),Cm(21)))
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    normal=doc.styles["Normal"]; normal.font.name=f["body_font"]; normal.font.size=Pt(f["body_font_size"]); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for n,s,c,b,a in (("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)):
        st=doc.styles[n]; st.font.name=f["heading_font"]; st.font.size=Pt(s); st.font.bold=True; st.font.color.rgb=c; st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a); st.paragraph_format.keep_with_next=True
    for n in ("List Bullet","List Number"):
        st=doc.styles[n]; st.font.name=f["body_font"]; st.font.size=Pt(f["body_font_size"]); st.paragraph_format.left_indent=Inches(.5); st.paragraph_format.first_line_indent=Inches(-.25)
    cap=doc.styles["Caption"]; cap.font.name=f["body_font"]; cap.font.size=Pt(9); cap.font.italic=True; cap.font.color.rgb=MUTED; cap.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    code=doc.styles.add_style("Code Block",WD_STYLE_TYPE.PARAGRAPH); code.font.name="Consolas"; code.font.size=Pt(8.5); code.paragraph_format.left_indent=Inches(.25)
    if f.get("include_header"):
        p=sec.header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.add_run(f"{r['title']} | {r.get('selected_version','')} | Generated {date}")
    if f.get("include_footer"):
        p=sec.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if f.get("include_page_numbers"): p.add_run("Page "); field(p,"PAGE","1")
    up=OxmlElement("w:updateFields"); up.set(qn("w:val"),"true"); doc.settings._element.append(up)

def heading(doc,n,title):
    if n>1: doc.add_page_break()
    doc.add_heading(f"{n}. {title}",1)

def margins(cell):
    pr=cell._tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar"); pr.append(mar)
    for edge,value in (("top",80),("bottom",80),("start",120),("end",120)):
        x=OxmlElement(f"w:{edge}"); x.set(qn("w:w"),str(value)); x.set(qn("w:type"),"dxa"); mar.append(x)

def table(doc,headers,rows,widths):
    if sum(widths)!=WIDTH: raise ReportError("Invalid table width")
    t=doc.add_table(rows=1,cols=len(headers),style="Table Grid"); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    pr=t._tbl.tblPr
    for tag,value in (("w:tblW",WIDTH),("w:tblInd",INDENT)):
        x=OxmlElement(tag); x.set(qn("w:w"),str(value)); x.set(qn("w:type"),"dxa"); pr.append(x)
    grid=t._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for value in widths: x=OxmlElement("w:gridCol"); x.set(qn("w:w"),str(value)); grid.append(x)
    for i,value in enumerate(headers):
        t.rows[0].cells[i].text=str(value); sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),"F2F4F7"); t.rows[0].cells[i]._tc.get_or_add_tcPr().append(sh)
        for run in t.rows[0].cells[i].paragraphs[0].runs: run.bold=True
    mark=OxmlElement("w:tblHeader"); mark.set(qn("w:val"),"true"); t.rows[0]._tr.get_or_add_trPr().append(mark)
    for values in rows:
        cells=t.add_row().cells
        for i,value in enumerate(values): cells[i].text="" if value is None else str(value)
    for row in t.rows:
        for i,cell in enumerate(row.cells):
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; margins(cell); tc=cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc.set(qn("w:w"),str(widths[i])); tc.set(qn("w:type"),"dxa")
    doc.add_paragraph(); return t

def expand(patterns):
    found=[]
    for pattern in patterns: found.extend(p for p in ROOT.glob(pattern) if p.is_file())
    return list(dict.fromkeys(sorted(found,key=lambda p:str(p).lower())))

def clean(text):
    text=re.sub(r"!\[[^]]*\]\([^)]+\)","",text); text=re.sub(r"\[([^]]+)\]\([^)]+\)",r"\1",text)
    return text.replace("**","").replace("__","").replace(chr(96),"").strip()

def select(path,topics,rules):
    lines=path.read_text(encoding="utf-8",errors="replace").splitlines(); excluded={x.lower() for x in rules.get("exclude_headings",[])}
    chunks=[]; current=[]; title=""
    for line in lines:
        m=re.match(r"^(#{1,6})\s+(.+)$",line)
        if m:
            if current: chunks.append((title,current))
            title=clean(m.group(2)); current=[line]
        else: current.append(line)
    if current: chunks.append((title,current))
    terms=[x.lower() for x in topics]; chosen=[]
    for title,chunk in chunks:
        if title.lower() not in excluded and (not terms or any(x in " ".join(chunk).lower() for x in terms)): chosen+=chunk+[""]
    if not chosen: chosen=lines[:80]
    limit=int(rules.get("maximum_paragraphs_per_source",8)); out=[]; count=0
    for line in chosen:
        if line.strip() and (not out or not out[-1].strip()): count+=1
        if count>limit: break
        out.append(line)
    return out

def render(doc,lines,seen,rules):
    i=0; buf=[]; code=[]; inside=False
    def flush():
        if buf:
            value=clean(" ".join(x.strip() for x in buf)); key=re.sub(r"\W+","",value.lower())
            if value and (not rules.get("avoid_duplicate_paragraphs") or key not in seen): doc.add_paragraph(value,"Normal"); seen.add(key)
            buf.clear()
    while i<len(lines):
        line=lines[i]; s=line.strip()
        if s.startswith(chr(96)*3):
            flush()
            if inside and code: doc.add_paragraph("\n".join(code),"Code Block"); code=[]
            inside=not inside; i+=1; continue
        if inside: code.append(line); i+=1; continue
        m=re.match(r"^(#{1,6})\s+(.+)$",line)
        if m: flush(); i+=1; continue
        if s.startswith("|") and i+1<len(lines) and re.match(r"^\|?[\s:|-]+\|?$",lines[i+1].strip()):
            flush(); raw=[s]; i+=2
            while i<len(lines) and lines[i].strip().startswith("|"): raw.append(lines[i].strip()); i+=1
            values=[[clean(c) for c in row.strip("|").split("|")] for row in raw]; n=len(values[0]); widths=[WIDTH//n]*n; widths[-1]+=WIDTH-sum(widths)
            table(doc,values[0],values[1:],widths); continue
        bullet=re.match(r"^\s*[-*+]\s+(.+)$",line); number=re.match(r"^\s*\d+[.)]\s+(.+)$",line)
        if bullet or number: flush(); doc.add_paragraph(clean((bullet or number).group(1)),style="List Bullet" if bullet else "List Number"); i+=1; continue
        if s: buf.append(line)
        else: flush()
        i+=1
    flush()

def sources(doc,patterns,topics,cfg,seen):
    files=expand(patterns)
    if not files and cfg["content_rules"].get("missing_source_behavior")!="continue": raise ReportError(f"No sources match {patterns}")
    limits=cfg["content_rules"].get("section_limits",{}).get("implementation_subsection",{})
    rules=dict(cfg["content_rules"]); rules["maximum_paragraphs_per_source"]=min(2,int(limits.get("maximum_paragraphs",3)))
    for path in files[:2]: render(doc,select(path,topics,rules),seen,rules)

def image(doc,path,caption,cfg):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    shape=p.add_run().add_picture(str(path),width=Inches(float(cfg["format"]["image_max_width_inches"])))
    shape._inline.docPr.set("descr",caption); shape._inline.docPr.set("title",caption); doc.add_paragraph(caption,"Caption"); return True

def find_screenshots(cfg):
    out=[]
    for directory in cfg.get("directories",[]):
        root=ROOT/directory
        if not root.exists(): continue
        files=root.rglob("*") if cfg.get("recursive") else root.glob("*")
        for path in files:
            if path.suffix.lower() not in {".png",".jpg",".jpeg"}: continue
            for mapping in cfg.get("mappings",[]):
                if fnmatch.fnmatch(path.name.lower(),mapping["pattern"].lower()): out.append((mapping["title"],path)); break
    return list(dict.fromkeys(out))

def find_eda(cfg):
    root=ROOT/cfg["root"]
    dates=sorted([p for p in root.glob(cfg["date_prefix"]+"*") if p.is_dir()],key=lambda p:p.name.removeprefix(cfg["date_prefix"]))
    if not dates:return None,[],{}
    hours=sorted([p for p in dates[-1].glob(cfg["hour_prefix"]+"*") if p.is_dir()],key=lambda p:p.name.removeprefix(cfg["hour_prefix"]))
    if not hours:return None,[],{}
    batches=sorted([p for p in hours[-1].glob(cfg["batch_prefix"]+"*") if p.is_dir()],key=lambda p:p.name.removeprefix(cfg["batch_prefix"]))
    if not batches:return None,[],{}
    batch=batches[-1]; images=[]
    for pattern in cfg.get("image_patterns",["*.png"]): images+=list(batch.glob(pattern))
    summary=batch/cfg.get("summary_file",""); data=json.loads(summary.read_text(encoding="utf-8")) if summary.is_file() else {}
    return batch.name.removeprefix(cfg["batch_prefix"]),sorted(set(images)),data

def find_features(cfg):
    db=next((ROOT/p for p in cfg.get("paths",[]) if (ROOT/p).is_file()),None)
    if not db:return None,[]
    result=[]
    with sqlite3.connect(db) as con:
        names=[r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
        for name in names:
            safe=name.replace('"','""'); info=con.execute(f'PRAGMA table_info("{safe}")').fetchall(); count=con.execute(f'SELECT COUNT(*) FROM "{safe}"').fetchone()[0]
            samples=[]
            if cfg.get("include_sample_rows"):
                cols=[x[1] for x in info]; samples=[dict(zip(cols,row)) for row in con.execute(f'SELECT * FROM "{safe}" LIMIT ?',(int(cfg.get("sample_row_limit",5)),))]
            result.append({"name":name,"rows":count,"columns":[f"{x[1]} ({x[2] or 'ANY'})" for x in info],"pk":[x[1] for x in info if x[5]],"samples":samples})
    return db,result

def find_dvc(cfg):
    path=ROOT/cfg["pipeline_file"]
    if not path.is_file(): return [], "", ""
    stages=[]
    for name,data in (yaml.safe_load(path.read_text(encoding="utf-8")) or {}).get("stages",{}).items():
        stages.append((name,data.get("cmd",""),"\n".join(map(str,data.get("deps",[]))),"\n".join(map(str,data.get("outs",[])))))
    status=dag=""
    if cfg.get("run_cli_when_available") and shutil.which("dvc"):
        for args,target in ((["dvc","status"],"status"),(["dvc","dag"],"dag")):
            try: value=(subprocess.run(args,cwd=ROOT,capture_output=True,text=True,timeout=30).stdout or "").strip()
            except Exception: value=""
            if target=="status": status=value
            else: dag=value
    return stages,status,dag

def find_mlflow(cfg):
    runs=[]
    for root in [ROOT/p for p in cfg.get("tracking_locations",[]) if (ROOT/p).is_dir()]:
        for meta in root.glob("*/*/meta.yaml"):
            data=yaml.safe_load(meta.read_text(encoding="utf-8")) or {}
            if cfg.get("latest_successful_runs") and str(data.get("status")) not in ("3","FINISHED"): continue
            exp_meta=meta.parents[1]/"meta.yaml"; exp=yaml.safe_load(exp_meta.read_text(encoding="utf-8")) if exp_meta.is_file() else {}
            if cfg.get("experiment_name") and exp.get("name")!=cfg["experiment_name"]: continue
            folder=meta.parent; params={p.name:p.read_text(encoding="utf-8") for p in (folder/"params").glob("*")} if (folder/"params").is_dir() else {}; metrics={}
            if (folder/"metrics").is_dir():
                for p in (folder/"metrics").glob("*"):
                    try: metrics[p.name]=float(p.read_text(encoding="utf-8").splitlines()[-1].split()[1])
                    except (ValueError,IndexError): pass
            runs.append({"experiment":exp.get("name",""),"name":data.get("run_name",folder.name),"run_id":data.get("run_id",folder.name),"status":"FINISHED","params":params,"metrics":metrics,"start":int(data.get("start_time",0))})
    runs.sort(key=lambda x:x["start"],reverse=True)
    return runs[:max(1,int(cfg.get("runs_per_model",1))*2)]

def add_features(doc,items):
    if not items: doc.add_paragraph("Evidence not available at report-generation time."); return
    table(doc,("Table","Rows","Columns and Datatypes","Primary Key"),((x["name"],x["rows"],", ".join(x["columns"]),", ".join(x["pk"]) or "None") for x in items),(1700,900,5100,1660))
    for item in sorted(items,key=lambda x:x["rows"],reverse=True)[:3]:
        if not item["samples"]: continue
        doc.add_heading(f"Sample: {item['name']}",3); keys=list(item["samples"][0])[:5]; widths=[WIDTH//len(keys)]*len(keys); widths[-1]+=WIDTH-sum(widths)
        table(doc,keys,([str(row.get(k,""))[:80] for k in keys] for row in item["samples"]),widths)

def add_dvc(doc,data):
    stages,status,dag=data
    if not stages: doc.add_paragraph("Evidence not available at report-generation time."); return
    table(doc,("Stage","Command","Dependencies","Outputs"),stages,(1300,2900,2600,2560))
    if status: doc.add_heading("DVC Status",3); doc.add_paragraph(status,"Code Block")
    if dag: doc.add_heading("DVC DAG",3); doc.add_paragraph(dag,"Code Block")

def add_mlflow(doc,runs,cfg):
    if not runs: doc.add_paragraph("Evidence not available at report-generation time."); return
    options=cfg["artifacts"]["mlflow"]
    for run in runs:
        doc.add_heading(f"{run['experiment']} - {run['name']}",3); rows=[]
        if options.get("include_run_ids"): rows.append(("Run ID",run["run_id"]))
        rows.append(("Status",run["status"]))
        if options.get("include_parameters"): rows.append(("Parameters",", ".join(f"{k}={v}" for k,v in run["params"].items()) or "None"))
        if options.get("include_metrics"): rows.append(("Metrics",", ".join(f"{k}={v:.6g}" for k,v in run["metrics"].items()) or "None"))
        table(doc,("Field","Value"),rows,(1800,7560))

def matches(patterns,latest=False,limit=None):
    files=expand(patterns)
    if latest and files: files=[max(files,key=lambda p:p.stat().st_mtime)]
    return files[:int(limit)] if limit else files

def hyperlink(paragraph,label,path):
    from docx.opc.constants import RELATIONSHIP_TYPE
    relative=Path(path).resolve().relative_to(ROOT).as_posix()
    rid=paragraph.part.relate_to(relative,RELATIONSHIP_TYPE.HYPERLINK,is_external=True)
    link=OxmlElement("w:hyperlink"); link.set(qn("r:id"),rid); run=OxmlElement("w:r"); props=OxmlElement("w:rPr")
    color=OxmlElement("w:color"); color.set(qn("w:val"),"0563C1"); underline=OxmlElement("w:u"); underline.set(qn("w:val"),"single")
    props.extend((color,underline)); run.append(props); text=OxmlElement("w:t"); text.text=f"{label} ({relative})"; run.append(text); link.append(run); paragraph._p.append(link)

def folder_text(root,max_depth):
    lines=[root.relative_to(ROOT).as_posix()+"/"]
    for path in sorted(root.rglob("*"),key=lambda p:str(p).lower()):
        depth=len(path.relative_to(root).parts)
        if depth<=int(max_depth): lines.append("  "*(depth-1)+("- " if path.is_file() else "+ ")+path.name)
    return "\n".join(lines)

def json_rows(value,prefix=""):
    rows=[]
    if isinstance(value,dict):
        for key,item in value.items(): rows+=json_rows(item,f"{prefix}.{key}".strip("."))
    elif isinstance(value,list): rows.append((prefix,json.dumps(value)[:500]))
    else: rows.append((prefix,value))
    return rows

def dataset_evidence(doc,spec):
    import csv
    files=matches(spec.get("source_patterns",[]),limit=spec.get("max_files")); rows=[]
    for path in files:
        try:
            if path.suffix.lower()==".csv":
                with path.open(encoding="utf-8",errors="replace",newline="") as handle:
                    reader=csv.reader(handle); header=next(reader,[]); sample=[]; count=0
                    for count,row in enumerate(reader,1):
                        if count<=int(spec.get("sample_row_limit",5)): sample.append(row)
                rows.append((path.relative_to(ROOT),count,len(header)))
                if spec.get("include_sample_rows") and sample and header:
                    widths=[WIDTH//len(header)]*len(header); widths[-1]+=WIDTH-sum(widths); table(doc,header,sample,widths)
            else: rows.append((path.relative_to(ROOT),"Available","Parquet"))
        except Exception: continue
    if rows: table(doc,("Dataset","Rows","Columns / Type"),rows,(5200,1600,2560)); return True
    return False

def evidence(doc,spec,cfg,seen,context):
    kind=spec.get("type"); found=False
    if kind=="document_link":
        for path in matches(spec.get("file_patterns",[]),spec.get("latest_only",False),1 if spec.get("latest_only") else None):
            hyperlink(doc.add_paragraph(),spec.get("label",path.name),path); found=True
    elif kind=="code_excerpt":
        for path in matches(spec.get("source_patterns",[]),limit=1):
            try:
                lines=path.read_text(encoding="utf-8",errors="replace").splitlines()[:min(18,int(spec.get("maximum_lines",spec.get("max_lines_per_file",18))))]
                doc.add_heading(path.relative_to(ROOT).as_posix(),3); doc.add_paragraph("\n".join(lines),"Code Block"); found=True
            except OSError: pass
    elif kind=="log_excerpt":
        for path in matches(spec.get("source_patterns",[]),spec.get("latest_only",False),spec.get("max_files")):
            lines=path.read_text(encoding="utf-8",errors="replace").splitlines()[-min(15,int(spec.get("maximum_lines",spec.get("max_lines_per_file",15)))):]
            doc.add_heading(path.relative_to(ROOT).as_posix(),3); doc.add_paragraph("\n".join(lines),"Code Block"); found=True
    elif kind=="screenshot":
        for path in matches(spec.get("file_patterns",[]),limit=min(3,int(spec.get("maximum_images",3)))):
            image(doc,path,path.stem.replace("_"," ").title(),cfg); found=True
    elif kind=="latest_eda_batch":
        batch,images,summary=context["eda"]
        if batch: doc.add_paragraph(f"Selected latest partitioned EDA batch: {batch}."); found=True
        if spec.get("include_summary_json") and summary: table(doc,("Measure","Value"),json_rows(summary),(3000,6360))
        if spec.get("include_all_images"):
            for path in images: image(doc,path,f"{path.stem.replace('_',' ').title()} - EDA batch {batch}",cfg)
    elif kind=="sqlite_summary":
        items=context["features"][1]; add_features(doc,items); found=bool(items)
    elif kind=="dvc_summary":
        stages,status,dag=context["dvc"]; add_dvc(doc,(stages,status,dag)); found=bool(stages)
        lock=ROOT/cfg["artifacts"]["dvc"]["lock_file"]
        if lock.is_file():
            locked=(yaml.safe_load(lock.read_text(encoding="utf-8")) or {}).get("stages",{})
            table(doc,("Locked Stage","Dependencies","Outputs"),((k,len(v.get("deps",[])),len(v.get("outs",[]))) for k,v in locked.items()),(4000,2680,2680)); found=True
    elif kind=="mlflow_summary":
        add_mlflow(doc,context["runs"],cfg); found=bool(context["runs"])
    elif kind=="markdown_summary":
        for path in expand(spec.get("sources",[]))[:1]:
            rules=dict(cfg["content_rules"]); rules["maximum_paragraphs_per_source"]=spec.get("maximum_paragraphs",4)
            render(doc,select(path,spec.get("include_topics",[]),rules),seen,rules); found=True
    elif kind=="folder_tree":
        for root in [ROOT/p for p in spec.get("root_patterns",[]) if (ROOT/p).exists()]:
            doc.add_paragraph(folder_text(root,spec.get("max_depth",3)),"Code Block"); found=True
    elif kind=="json_summary":
        for path in matches(spec.get("file_patterns",[]),spec.get("latest_only",False),1 if spec.get("latest_only") else None):
            try: table(doc,("Measure","Value"),json_rows(json.loads(path.read_text(encoding="utf-8"))),(3000,6360)); found=True
            except Exception: pass
    elif kind=="dataset_summary": found=dataset_evidence(doc,spec)
    elif kind=="feature_logic_table":
        rows=[(p.name,p.suffix.lstrip(".").upper(),p.relative_to(ROOT)) for p in matches(spec.get("source_patterns",[]))[:10]]
        if rows: table(doc,("Source","Type","Location"),rows,(2600,1200,5560)); found=True
    if not found: doc.add_paragraph("Evidence not available at report-generation time.")

def add_deliverables(doc,section,cfg,seen,context,number):
    block=cfg.get(section.get("evidence_configuration",""),{}); deliverables=block.get("deliverables",[])
    if len(deliverables)!=10: raise ReportError("Results section requires exactly ten configured deliverables")
    for item in deliverables:
        doc.add_heading(f"{number}.{item['id']} {item['title']}",2)
        for requirement in item.get("requirements",[]): doc.add_paragraph(requirement,"List Bullet")
        for spec in item.get("evidence",[]): evidence(doc,spec,cfg,seen,context)
def add_inline_markdown(paragraph,text):
    parts=re.split(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*)",text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"): paragraph.add_run(part[2:-2]).bold=True
        elif part.startswith("*") and part.endswith("*"): paragraph.add_run(part[1:-1]).italic=True
        else: paragraph.add_run(part)

def render_markdown_file(doc,section):
    path=(ROOT/section["source_file"]).resolve()
    validation=section.get("validation",{}); rendering=section.get("rendering",{})
    if validation.get("source_required") and not path.is_file(): raise ReportError(f"Required source missing: {path}")
    raw=path.read_text(encoding="utf-8",errors="strict")
    if validation.get("reject_empty_content") and not raw.strip(): raise ReportError("Problem Statement source is empty")
    if validation.get("reject_tables") and any(line.strip().startswith("|") for line in raw.splitlines()): raise ReportError("Problem Statement source contains a table")
    if validation.get("reject_code_blocks") and chr(96)*3 in raw: raise ReportError("Problem Statement source contains a code block")
    paragraphs=[]; current=[]
    for line in raw.splitlines():
        stripped=line.strip()
        if stripped.startswith("#"):
            if current: paragraphs.append(" ".join(current)); current=[]
            if rendering.get("include_source_heading") and stripped.startswith("# "): paragraphs.append(stripped[2:].strip())
            continue
        if re.match(r"^\s*(?:[-*+] |\d+[.)] )",line):
            if current: paragraphs.append(" ".join(current)); current=[]
            if rendering.get("render_lists"): paragraphs.append(re.sub(r"^\s*(?:[-*+] |\d+[.)] )","",line))
            continue
        if not stripped:
            if current: paragraphs.append(" ".join(current)); current=[]
        elif rendering.get("render_paragraphs",True): current.append(stripped)
    if current: paragraphs.append(" ".join(current))
    maximum=int(rendering.get("maximum_words",500)); total=sum(len(p.split()) for p in paragraphs)
    if total>maximum: raise ReportError(f"Problem Statement exceeds {maximum} words")
    for value in paragraphs:
        paragraph=doc.add_paragraph()
        if rendering.get("preserve_inline_formatting"): add_inline_markdown(paragraph,value)
        else: paragraph.add_run(clean(value))
def owned_sources(section,cfg,required):
    if section.get("source")!=required: raise ReportError(f"{section['id']} cannot use {section.get('source')}")
    return cfg["content_sources"][required]

REJECTED_CONCEPTS=("repository structure","folder tree","transition matrix","install","coding standard","contributor","configuration reference","sql schema")

def knowledge_sentences(patterns,topics,maximum_words=300):
    candidates=[]
    for path in expand(patterns)[:12]:
        text=path.read_text(encoding="utf-8",errors="replace")
        text=re.sub("(?ms)"+chr(96)*3+".*?"+chr(96)*3," ",text)
        for line in text.splitlines():
            value=clean(re.sub(r"^\s*(?:#{1,6}|[-*+] |\d+[.)] )\s*","",line)).strip()
            low=value.lower()
            if len(value)<35 or value.startswith("|") or any(x in low for x in REJECTED_CONCEPTS): continue
            score=sum(1 for topic in topics if topic.lower() in low)
            if score: candidates.append((score,value))
    candidates.sort(key=lambda x:-x[0]); chosen=[]; seen=set(); words=0
    for _,value in candidates:
        for sentence in re.split(r"(?<=[.!?])\s+",value):
            key=re.sub(r"\W+","",sentence.lower()); count=len(sentence.split())
            if key and key not in seen and 8<=count<=55 and words+count<=maximum_words:
                chosen.append(sentence); seen.add(key); words+=count
    return chosen

def prose(doc,sentences,max_paragraphs=4):
    if not sentences: doc.add_paragraph("Evidence not available at report-generation time."); return
    size=max(1,(len(sentences)+max_paragraphs-1)//max_paragraphs)
    for index in range(0,len(sentences),size): doc.add_paragraph(" ".join(sentences[index:index+size]))

def semantic_problem(doc,section,cfg):
    topics=("recommendation","business","heterogeneous","data quality","reproducibility","feature consistency","tracking","orchestration","challenge","motivation")
    prose(doc,knowledge_sentences(owned_sources(section,cfg,"assignment"),topics,340),4)

def semantic_objectives(doc,section,cfg):
    sentences=knowledge_sentences(owned_sources(section,cfg,"assignment"),("objective","goal","deliverable","pipeline","validation","feature","model","tracking","orchestration"),240)
    for sentence in sentences[:10]: doc.add_paragraph(sentence,"List Number")

def render_pipeline_image(document: Document,repository_root: Path,image_config: dict) -> None:
    import hashlib
    relative=Path(image_config["path"]); path=(repository_root/relative).resolve()
    if image_config.get("required") and not path.is_file():
        raise ReportError(f"Required pipeline image not found:\n{relative.as_posix()}")
    if not path.is_file(): return
    digest=hashlib.sha256(path.read_bytes()).hexdigest()
    if digest in _IMAGE_HASHES: return
    metadata=DocxImage.from_file(str(path)); width,height=int(metadata.width),int(metadata.height)
    max_width=int(Inches(float(image_config.get("max_width_inches",6.3))))
    max_height=int(Inches(float(image_config.get("max_height_inches",7.5))))
    scale=min(max_width/width,max_height/height,1.0)
    paragraph=document.add_paragraph()
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER if image_config.get("alignment")=="center" else WD_ALIGN_PARAGRAPH.LEFT
    shape=paragraph.add_run().add_picture(str(path),width=Emu(int(width*scale)),height=Emu(int(height*scale)))
    caption=image_config.get("caption","")
    shape._inline.docPr.set("descr",caption); shape._inline.docPr.set("title",caption)
    _IMAGE_HASHES.add(digest)
    if caption:
        caption_paragraph=document.add_paragraph(caption,"Caption"); caption_paragraph.alignment=paragraph.alignment

def semantic_methodology(doc,section,cfg):
    rendering=section.get("rendering",{})
    if rendering.get("generate_pipeline_image"): raise ReportError("Generated pipeline diagrams are disabled")
    if rendering.get("include_transition_table"): raise ReportError("Methodology transition tables are disabled")
    if rendering.get("include_overview"):
        sentences=knowledge_sentences(section.get("sources",[]),("pipeline","architecture","flow","stage"),160)
        prose(doc,sentences,int(rendering.get("maximum_overview_paragraphs",2)))
    if rendering.get("include_pipeline_image"): render_pipeline_image(doc,ROOT,section["pipeline_image"])
    if rendering.get("include_stage_summary"):
        maximum=int(rendering.get("maximum_stage_words",60))
        for stage in section.get("pipeline_stages",[]):
            doc.add_heading(stage["name"],3)
            words=stage["description"].split()
            doc.add_paragraph(" ".join(words[:maximum]))
def semantic_implementation(doc,section):
    for index,sub in enumerate(section.get("subsections",[]),1):
        doc.add_heading(f"6.{index} {sub['title']}",2)
        sentences=knowledge_sentences(sub.get("sources",[]),(sub["title"].lower(),"purpose","input","output","technology","implementation"),180); joined=" ".join(sentences)
        fields=(("Purpose",sentences[0] if sentences else f"Supports {sub['title'].lower()} in the RecoMart lifecycle."),
                ("Implementation Summary"," ".join(sentences[1:3]) if len(sentences)>1 else "Implemented as a modular pipeline component using repository-defined contracts."),
                ("Technology Used",next((x for x in ("Python","SQLite","DVC","MLflow","Apache Airflow","Pandas","Parquet") if x.lower() in joined.lower()),"Python")),
                ("Input",next((x for x in sentences if "input" in x.lower()),"Validated upstream project data and configuration.")),
                ("Output",next((x for x in sentences if "output" in x.lower()),"Versioned downstream artifacts for the next pipeline stage.")))
        for label,value in fields:
            p=doc.add_paragraph(); r=p.add_run(label+": "); r.bold=True; p.add_run(value)

def semantic_conclusion(doc,section):
    for title,topics in (("Project Achievements",("achievement","complete","implemented","reproducible")),("Limitations",("limitation","out of scope","constraint")),("Future Work",("future","enhancement","deployment","monitoring"))):
        doc.add_heading(title,2); prose(doc,knowledge_sentences(section.get("sources",[]),topics,95),1)
def select_output(cfg):
    v=cfg["report"]["versioning"]; directory=ROOT/v["output_directory"]; directory.mkdir(parents=True,exist_ok=True)
    pattern=v["filename_template"].replace("{version}",v["prefix"]+"*"); numbers=[]
    for path in directory.glob(pattern):
        match=re.search(r"_v(\d+)\.docx$",path.name,re.I)
        if match: numbers.append(int(match.group(1)))
    number=max(numbers,default=int(v.get("starting_version",1))-1)+1; version=f"{v['prefix']}{number}"
    return directory/v["filename_template"].format(version=version),version,set(directory.glob(pattern))

def generate(config_path=DEFAULT_CONFIG,overwrite=False):
    cfg=load_config(config_path.resolve()); _IMAGE_HASHES.clear(); output,version,existing=select_output(cfg)
    if output.exists(): raise FileExistsError(f"Refusing to overwrite {output}")
    date=datetime.now().astimezone().strftime("%Y-%m-%d"); doc=Document(); cfg["report"]["selected_version"]=version; setup(doc,cfg,date); seen=set()
    shots=find_screenshots(cfg["artifacts"]["screenshots"]); batch,eda_images,eda_summary=find_eda(cfg["artifacts"]["eda"])
    feature_db,features=find_features(cfg["artifacts"]["feature_store"]); dvc=find_dvc(cfg["artifacts"]["dvc"]); runs=find_mlflow(cfg["artifacts"]["mlflow"])
    enabled=[s for s in cfg["sections"] if s.get("enabled",True)]
    for number,section in enumerate(enabled,1):
        heading(doc,number,section["title"]); mode=section.get("render_mode",section["type"])
        if mode=="title_page":
            for _ in range(4): doc.add_paragraph()
            for value,size,color,bold in ((cfg["report"]["institution"],12,MUTED,True),(cfg["report"]["title"],26,DARK,True),(cfg["report"]["subtitle"],15,BLUE,False)):
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(value); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
            values=(cfg["report"]["assignment"],f"Report Version: {version}")
            if cfg["report"].get("generated_date"): values+=(f"Generated: {date}",)
            for value in values: p=doc.add_paragraph(value); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif mode=="team_table":
            table(doc,("S. No.","Name","BITS ID / Email ID"),((m["serial_no"],m["name"],m["bits_id"]) for m in cfg["team"]),(1000,3300,5060))
            if cfg["format"].get("include_table_of_contents"):
                p=doc.add_paragraph(); r=p.add_run("Table of Contents"); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=BLUE
                toc=doc.add_paragraph(); field(toc,r'TOC \o "1-3" \h \z \u',"Update field in Word to refresh.")
        elif mode=="markdown_file": render_markdown_file(doc,section)
        elif mode in ("executive_summary","narrative"): semantic_problem(doc,section,cfg)
        elif mode in ("numbered_objectives","numbered_list"): semantic_objectives(doc,section,cfg)
        elif mode=="methodology": semantic_methodology(doc,section,cfg)
        elif mode=="implementation": semantic_implementation(doc,section)
        elif mode=="evidence":
            context={"eda":(batch,eda_images,eda_summary),"features":(feature_db,features),"dvc":dvc,"runs":runs}
            add_deliverables(doc,section,cfg,seen,context,number)
        elif mode=="conclusion": semantic_conclusion(doc,section)
        else: raise ReportError(f"Unsupported render mode: {mode}")
    seen_text=set()
    for paragraph in list(doc.paragraphs):
        value=paragraph.text.strip(); key=re.sub(r"\\W+","",value.lower())
        if value and not paragraph.style.name.startswith("Heading") and paragraph.style.name!="Caption":
            if key in seen_text: paragraph._element.getparent().remove(paragraph._element)
            else: seen_text.add(key)
    output.parent.mkdir(parents=True,exist_ok=True); doc.save(output); validate(output,cfg,version,existing)
    return output,{"eda_batch":batch,"screenshots":[p.name for _,p in shots],"feature_tables":[x["name"] for x in features],"dvc_stages":[x[0] for x in dvc[0]],"mlflow_runs":[x["run_id"] for x in runs]}

def validate(path,cfg,version,existing):
    if path.name!=cfg["report"]["versioning"]["filename_template"].format(version=version): raise ReportError("Generated filename version mismatch")
    if any(not p.exists() for p in existing): raise ReportError("An earlier report was overwritten")
    if cfg["validation"].get("require_output_file") and (not path.is_file() or path.stat().st_size==0): raise ReportError("Output missing or empty")
    doc=Document(path); actual=[p.text for p in doc.paragraphs if p.style.name=="Heading 1"]; expected=[f"{i}. {s['title']}" for i,s in enumerate([x for x in cfg["sections"] if x.get("enabled",True)],1)]
    if cfg["validation"].get("reject_additional_top_level_sections") and actual!=expected: raise ReportError("Top-level headings do not match YAML")
    deliverables=cfg.get("results_and_outputs",{}).get("deliverables",[])
    configured=[f"7.{item['id']} {item['title']}" for item in deliverables]
    headings2=[p.text for p in doc.paragraphs if p.style.name=="Heading 2"]
    if len(deliverables)!=10 or any(title not in headings2 for title in configured): raise ReportError("Configured results deliverables are missing or out of structure")
    if cfg["validation"].get("require_all_team_members"):
        text="\n".join(p.text for p in doc.paragraphs)+"\n"+"\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
        if any(m["name"] not in text or m["bits_id"] not in text for m in cfg["team"]): raise ReportError("A configured team member is absent")
    if any(x in "\n".join(p.text for p in doc.paragraphs) for x in ("{{","}}","<PLACEHOLDER>")): raise ReportError("Unresolved template marker")

def main(argv:Sequence[str]|None=None):
    parser=argparse.ArgumentParser(); parser.add_argument("--config",type=Path,default=DEFAULT_CONFIG); parser.add_argument("--overwrite",action="store_true"); args=parser.parse_args(argv)
    output,evidence=generate(args.config,args.overwrite); print(json.dumps({"output":str(output),**evidence},indent=2)); return 0
if __name__=="__main__": raise SystemExit(main())
